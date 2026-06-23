"""Generate Word/docx files from the same Markdown used by report preview."""

from __future__ import annotations

import io
import re
from collections.abc import Mapping
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.utils.report_sections import generate_formal_report


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
FALLBACK_REPORT = """# 实验报告生成提示

当前报告内容为空或上下文不足。

请先完成虚拟实验流程，或在报告生成页面选择样本、场景和计算结果。缺失数据将显示为“待补充/未完成”，不会伪造文献或工程结论。

本文件仍可正常打开，用于提示 docx 导出链路没有崩溃。
"""


def _clean_text(value: object, fallback: str = "待补充") -> str:
    """Return text safe for python-docx cells and paragraphs."""

    if value is None:
        return fallback
    text = str(value)
    if not text.strip():
        return fallback
    return text.replace("\x00", "").strip()


def _set_run_font(run: Any, *, east_asia: str = "Microsoft YaHei", size: int = 11) -> None:
    run.font.name = east_asia
    run.font.size = Pt(size)
    if run.element.rPr is not None:
        run.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def _add_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(_clean_text(text))
    _set_run_font(run)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_clean_text(text))
    run.bold = True
    _set_run_font(run, east_asia="Microsoft YaHei", size=20)
    paragraph.paragraph_format.space_after = Pt(12)


def _add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(_clean_text(text), level=level)
    for run in heading.runs:
        run.bold = True
        _set_run_font(run, east_asia="Microsoft YaHei", size=16 if level == 1 else 13)


def _is_separator_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _parse_markdown_table(table_lines: list[str]) -> tuple[list[str], list[list[str]]]:
    """Parse simple GitHub-style pipe tables defensively."""

    parsed: list[list[str]] = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [_clean_text(cell, "") for cell in stripped.strip("|").split("|")]
        if _is_separator_row(cells):
            continue
        parsed.append(cells)

    if not parsed:
        return [], []

    headers = parsed[0]
    rows = parsed[1:] or [["待补充" for _ in headers]]
    width = max(len(headers), *(len(row) for row in rows))
    headers = headers + [""] * (width - len(headers))
    normalized_rows = []
    for row in rows:
        normalized_rows.append(row[:width] + [""] * max(0, width - len(row)))
    return headers, normalized_rows


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        _add_paragraph(doc, "表格内容待补充。")
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = _clean_text(header, "")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                _set_run_font(run, size=10)

    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = _clean_text(value, "")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, size=10)

    doc.add_paragraph()


def _markdown_from_input(report: Mapping[str, Any] | str | None) -> str:
    if isinstance(report, str):
        markdown = report
    elif isinstance(report, Mapping):
        try:
            markdown = generate_formal_report(dict(report))
        except Exception as exc:
            markdown = (
                "# 实验报告生成提示\n\n"
                f"报告上下文转换为 Markdown 时出现异常：{exc}\n\n"
                "请检查报告上下文是否包含非字符串、非数值或非表格兼容字段。"
            )
    elif report is None:
        markdown = FALLBACK_REPORT
    else:
        markdown = f"# 实验报告生成提示\n\n不支持的输入类型：{type(report).__name__}。"

    return markdown if markdown.strip() else FALLBACK_REPORT


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    """Convert a simple Markdown report to docx bytes."""

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)

    lines = _markdown_from_input(markdown_text).splitlines()
    idx = 0
    while idx < len(lines):
        raw_line = lines[idx]
        line = raw_line.strip()
        if not line:
            idx += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            headers, rows = _parse_markdown_table(table_lines)
            _add_table(doc, headers, rows)
            continue

        if line.startswith("# "):
            _add_title(doc, line[2:])
        elif line.startswith("## "):
            _add_heading(doc, line[3:], level=1)
        elif line.startswith("### "):
            _add_heading(doc, line[4:], level=2)
        elif line.startswith("- "):
            _add_paragraph(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            _add_paragraph(doc, line, style="List Number")
        elif line in {"$$", "---"}:
            pass
        else:
            _add_paragraph(doc, line)
        idx += 1

    if not doc.paragraphs and not doc.tables:
        _add_title(doc, "实验报告生成提示")
        _add_paragraph(doc, "当前报告内容为空。")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_docx(ctx: Mapping[str, Any] | str | None) -> bytes:
    """Generate a docx file as bytes.

    Args:
        ctx: Report context dict, Markdown string, or None.

    Returns:
        Non-empty Word document bytes suitable for st.download_button.
    """

    markdown = _markdown_from_input(ctx)
    return markdown_to_docx_bytes(markdown)
