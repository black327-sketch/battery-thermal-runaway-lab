from app.utils.report_generator import generate_report
from app.utils.report_context import collect_report_context
from docx import Document

from app.utils.report_docx import generate_docx, markdown_to_docx_bytes
from app.utils.report_sections import generate_formal_report, generate_formal_report_html


def test_report_contains_interactive_data_sources_and_disclaimer():
    report = generate_report(
        experiment_params={
            "interactive_state": {
                "selected_soc": 100,
                "leak_test_passed": True,
                "replacement_count": 3,
                "current_state": "report_generated",
                "gas_bag_filled": True,
                "gc_finished": True,
                "operation_logs": [{"time": "00:00:01", "action": "start_gc", "level": "info", "message": "GC 教学分析已启动。"}],
            },
            "gas_volume_status": {"status": "pending_user_input", "message": "待补充文献产气量计算公式和参数"},
            "score_summary": {"final_score": 100, "error_count": 0, "completion_pct": 100.0},
        },
        literature_data={"sample_info": {}, "gas_composition": {}, "flammable_composition": {}},
        calculation_results={"risk_info": {"level": "无法评价", "description": "测试"}},
    )
    assert "二维交互实验流程记录" in report
    assert "文献数据接入状态" in report
    assert "数据来源说明" in report
    assert "待补充文献项" in report
    assert "待补充文献产气量计算公式和参数" in report
    assert "本报告仅用于虚拟仿真教学" in report



def test_formal_report_html_and_docx_share_markdown_source():
    ctx = collect_report_context(
        experiment_params={"total_gas_l": 10.0, "scene_info": {"room_volume_m3": 10.0}},
        literature_data={
            "sample_info": {"sample_label": "demo"},
            "gas_composition": {"H2": 60.0, "CO2": 30.0, "CO": 10.0},
            "flammable_composition": {"H2": 60.0, "CO": 10.0},
        },
        calculation_results={
            "normalized": {"H2": 0.86, "CO": 0.14},
            "lfl_constants": {"H2": 4.0, "CO": 12.5},
            "lfl_mix": 4.5,
            "space_concentration": 0.1,
            "risk_ratio": 0.022,
            "risk_info": {"level": "\u6559\u5b66\u6a21\u578b\u4e0b\u4f4e\u98ce\u9669", "description": "test"},
        },
    )

    markdown = generate_formal_report(ctx)
    html = generate_formal_report_html(markdown)
    docx = generate_docx(ctx)

    assert markdown.startswith("# ")
    assert html.startswith("<!doctype html>")
    assert "<table>" in html
    assert "LFL_mix" in html
    assert len(docx) > 1000


def test_docx_generation_handles_markdown_tables_and_symbols(tmp_path):
    markdown = """# 测试报告

## 一、实验说明

本报告包含 H₂、CO₂、CH₄、C₂H₄、C₂H₆、LFL_mix、℃、ppm 等符号。

- 列表项一
- 列表项二

| 气体 | 含量 | 说明 |
| --- | --- | --- |
| H₂ | 20 ppm | 主要气体 |
| CO₂ | 30% | 非可燃组分 |
| CH₄ |  | 待补充 |
"""
    docx_bytes = markdown_to_docx_bytes(markdown)
    assert len(docx_bytes) > 1000

    out = tmp_path / "report.docx"
    out.write_bytes(docx_bytes)
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "测试报告" in text
    assert "H₂" in text
    assert doc.tables
    assert doc.tables[0].rows[1].cells[0].text == "H₂"


def test_docx_generation_returns_fallback_file_for_empty_report(tmp_path):
    docx_bytes = generate_docx("")
    out = tmp_path / "empty.docx"
    out.write_bytes(docx_bytes)

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "实验报告生成提示" in text
    assert "当前报告内容为空" in text
